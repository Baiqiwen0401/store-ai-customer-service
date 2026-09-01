from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_LINE_SPACING
from pathlib import Path

OUT = Path('个体门店AI客服技术方案_v1.0.docx')

NAVY = '0B2545'
BLUE = '2E74B5'
DARK_BLUE = '1F4D78'
MUTED = '5B6573'
LIGHT_BLUE = 'E8EEF5'
LIGHT_GRAY = 'F2F4F7'
CALLOUT = 'F4F6F9'
GOLD = '7A5A00'
RED = '9B1C1C'
WHITE = 'FFFFFF'
FONT = 'Microsoft YaHei'


def set_run_font(run, size=11, bold=False, color='000000', italic=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for side, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{side}'))
        if node is None:
            node = OxmlElement(f'w:{side}')
            tcMar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(dxa))
    tc_w.set(qn('w:type'), 'dxa')


def configure_table(table, widths, header=True, header_fill=LIGHT_GRAY):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.first_child_found_in('w:tblW')
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), '9360')
    tblW.set(qn('w:type'), 'dxa')
    tblInd = tblPr.first_child_found_in('w:tblInd')
    if tblInd is None:
        tblInd = OxmlElement('w:tblInd')
        tblPr.append(tblInd)
    tblInd.set(qn('w:w'), '120')
    tblInd.set(qn('w:type'), 'dxa')
    layout = tblPr.first_child_found_in('w:tblLayout')
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tblPr.append(layout)
    layout.set(qn('w:type'), 'fixed')
    grid = tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn('w:w'), str(width))
    for r_idx, row in enumerate(table.rows):
        if header and r_idx == 0:
            tr_pr = row._tr.get_or_add_trPr()
            tbl_header = OxmlElement('w:tblHeader')
            tbl_header.set(qn('w:val'), 'true')
            tr_pr.append(tbl_header)
        for c_idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[c_idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if header and r_idx == 0:
                shade_cell(cell, header_fill)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.1
                for run in p.runs:
                    set_run_font(run, size=9.5, bold=(header and r_idx == 0), color=NAVY if header and r_idx == 0 else '000000')


def set_paragraph_border(paragraph, color='D7DBE2', size='6', space='1'):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), space)
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def keep_with_next(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    node = OxmlElement('w:keepNext')
    pPr.append(node)


def add_page_field(paragraph):
    run = paragraph.add_run('第 ')
    set_run_font(run, size=8.5, color=MUTED)
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    paragraph._p.append(fld)
    run = paragraph.add_run(' 页')
    set_run_font(run, size=8.5, color=MUTED)


def add_text(doc, text, size=11, bold=False, color='000000', after=8, before=0, align=None, italic=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.333
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    r = p.add_run(text)
    return p


def add_callout(doc, title, body, color=CALLOUT, label_color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    configure_table(table, [9360], header=False)
    cell = table.cell(0, 0)
    shade_cell(cell, color)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=label_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        p = table.cell(0, i).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=9.5, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(value)
            set_run_font(r, size=9.5)
    configure_table(table, widths, header=True, header_fill=header_fill)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_flow(doc, items):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.2
    for idx, item in enumerate(items):
        r = p.add_run(item)
        set_run_font(r, size=10.5, bold=True, color=NAVY)
        if idx < len(items) - 1:
            sep = p.add_run('  ->  ')
            set_run_font(sep, size=10.5, color=BLUE)
    return p


def setup_styles(doc):
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    for level, size, color, before, after in [(1, 16, BLUE, 18, 10), (2, 13, BLUE, 12, 6), (3, 12, DARK_BLUE, 8, 4)]:
        style = doc.styles[f'Heading {level}']
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ['List Bullet', 'List Number']:
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('个体门店 AI 客服技术方案')
    set_run_font(r, size=8.5, color=MUTED)
    set_paragraph_border(p, color='D7DBE2', size='4')
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run('内部方案 | v1.0 | ')
    set_run_font(r, size=8.5, color=MUTED)
    add_page_field(p)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    setup_styles(doc)
    add_header_footer(doc)

    # Cover: customer_pack inspired title block.
    add_text(doc, '个体门店 AI 客服', size=28, bold=True, color=NAVY, after=4)
    add_text(doc, '一周试点上线技术方案', size=17, color=DARK_BLUE, after=16)
    lead = add_text(doc, '面向美容、美甲、宠物、健身等个体门店的可配置、可记忆、可人工接管的 AI 前台客服', size=11.5, color=MUTED, after=18)
    set_paragraph_border(lead, color='D7DBE2', size='6')
    add_table(doc, ['文档信息', '内容'], [
        ('方案定位', '首家试点门店 MVP；一周内可部署、可试运营'),
        ('首期渠道', '网页客服为必交付；企业微信为具备资质后的可选接入'),
        ('核心能力', '知识库问答、客户长期记忆、预约意向、转人工、店主纠错闭环'),
        ('版本', 'v1.0'),
        ('日期', '2026 年 9 月'),
    ], [2700, 6660], header_fill=LIGHT_BLUE)
    add_callout(doc, '方案结论', '一周内可以交付一个真实门店可使用的 AI 客服试点。首期以网页入口和标准咨询为主，企业微信、小程序、订单和发货按平台资质与验证结果逐步接入。')
    doc.add_page_break()

    add_heading(doc, '1. 项目目标与边界', 1)
    add_text(doc, '本项目建设一套面向个体门店的 AI 前台客服。系统以门店自有知识、客户档案和受控业务流程为依据，全天候处理高频咨询，沉淀客户关系，并在风险或不确定场景下将会话转交人工。')
    add_heading(doc, '1.1 首期业务目标', 2)
    for item in [
        '降低门店对营业时间、地址、价格、服务流程、注意事项等重复咨询的人工响应压力。',
        '在客户离开聊天前完成关键需求收集，并生成预约或跟进待办。',
        '将客户偏好、服务历史和已确认沟通结论沉淀为可审查的长期记忆。',
        '通过店主纠错和高质量案例审核，使客服内容持续贴合门店实际规则。',
    ]:
        add_bullet(doc, item)
    add_heading(doc, '1.2 首期范围', 2)
    add_table(doc, ['范围', '纳入首期', '暂不纳入首期'], [
        ('渠道', '独立网页咨询页与嵌入式聊天组件；企业微信接入预留', '小红书全自动私信、个人微信自动化、多平台统一收件箱'),
        ('客服', '标准问答、项目推荐、预约意向、人工接管', '医疗诊断、疗效保证、投诉闭环与退款裁决'),
        ('业务动作', '生成预约待办、记录客户档案、通知店员', '自动支付、自动退款、自动改价、自动发货'),
        ('学习机制', '记忆提取、店主纠错、审核后知识沉淀', '未经审核的模型自训练或将所有聊天直接写入知识库'),
    ], [1500, 3930, 3930])
    add_callout(doc, '核心原则', '模型负责理解、检索和建议；规则服务负责校验；业务系统负责真实执行。AI 不直接拥有数据库写权限或资金、退款、发货等高风险权限。', color='FFF8E8', label_color=GOLD)

    add_heading(doc, '2. 客户交互与服务流程', 1)
    add_text(doc, '网页客服是嵌入门店官网、预约页面或独立咨询页的聊天窗口。门店可通过二维码、短链接、公众号菜单、视频号或合规落地页把客户引导到该页面。')
    add_flow(doc, ['客户打开咨询页', 'AI 对话与检索', '回答或收集信息', '预约待办/人工接管'])
    add_heading(doc, '2.1 标准咨询示例', 2)
    add_table(doc, ['角色', '交互内容'], [
        ('客户', '敏感肌可以做深层清洁吗？周六下午有时间吗？'),
        ('AI 客服', '先提示需评估的皮肤状态，询问客户主要困扰；基于门店排班查询可预约时段并说明服务时长、价格与确认规则。'),
        ('客户', '主要是黑头，想约周六 16:30。'),
        ('AI 客服', '收集姓名、联系方式、服务项目和时间偏好，创建“待确认预约”并提示门店工作人员确认。'),
        ('店员', '在后台查看客户摘要、历史记录和预约信息，确认或人工跟进。'),
    ], [1450, 7910])
    add_heading(doc, '2.2 身份识别与跨会话记忆', 2)
    add_text(doc, '首次咨询可匿名进行；当客户预约、领取权益或希望保存历史时，使用手机号验证码或会员 ID 完成身份绑定。系统以客户主键而非浏览器 Cookie 关联跨设备、跨会话信息，店主可在后台查看、编辑和删除记忆。')

    add_heading(doc, '3. 总体技术架构', 1)
    add_flow(doc, ['网页/企业微信入口', '接入与会话服务', 'AI 编排服务', '知识/记忆/业务工具', '店主后台与人工坐席'])
    add_table(doc, ['层级', '职责', '首期实现建议'], [
        ('渠道接入层', '接收客户消息、身份校验、发送回复、会话状态管理', '网页聊天组件 + REST/WebSocket；企业微信作为适配器预留'),
        ('AI 编排层', '识别意图、检索资料、选择工具、生成回复、置信度判断', 'Dify/FastGPT 快速编排，或 FastAPI 服务封装模型与流程'),
        ('知识层', '检索门店资料、价目表、项目说明、服务政策与禁答规则', '文档解析 + 分段检索；门店级隔离的向量索引'),
        ('记忆与 CRM 层', '客户档案、历史对话、已确认记忆、预约与跟进记录', 'PostgreSQL + pgvector；对象按 tenant_id 隔离'),
        ('工具与规则层', '查询档期、创建待办、转人工、发送通知、权限校验', '受控 API 工具；每个动作记录审计日志'),
        ('运营后台', '门店配置、资料上传、会话处理、记忆管理、质检与数据看板', 'React/Next.js 管理后台或 Dify/FastGPT 运营界面加轻量业务后台'),
    ], [1600, 3400, 4360])
    add_heading(doc, '3.1 推荐部署形态', 2)
    add_text(doc, '首期采用 Docker Compose 在云服务器部署：Nginx 负责 HTTPS 与反向代理，应用服务处理会话和业务规则，PostgreSQL 保存业务数据与记忆，Redis 用于异步任务和限流，对象存储保存门店资料。模型调用使用可切换的国内模型 API，并保留 OpenAI 兼容接口。')

    add_heading(doc, '4. 长期记忆与持续专业化', 1)
    add_text(doc, '长期记忆不是完整聊天记录的无差别堆积，而是从对话中抽取、分类、审核并可撤销的业务事实。系统必须区分客户个体信息与门店通用知识，避免将某个客户的特殊情况错误推广。')
    add_table(doc, ['记忆类型', '典型内容', '写入方式', '调用规则'], [
        ('工作记忆', '本轮咨询目标、已问信息、当前预约进度', '自动写入，随会话过期', '仅当前会话使用'),
        ('客户记忆', '肤质偏好、预算、常用时段、已确认服务历史', 'AI 提取候选项；用户输入或店员确认后入库', '同一客户下次咨询时按相关性检索'),
        ('门店知识', '价目表、项目说明、营业时间、服务流程、禁答规则', '店主上传或编辑；发布后版本化', '所有该门店会话使用，优先级高于模型常识'),
        ('已验证案例', '店主修正过的高质量回答与常见问答', '人工审核后发布', '作为补充检索材料，并可生成 FAQ 草案'),
    ], [1500, 2770, 2640, 2450])
    add_heading(doc, '4.1 记忆闭环', 2)
    add_flow(doc, ['客户对话', '提取候选记忆', '规则去重与敏感检查', '店主确认/修正', '检索增强下次回答'])
    add_callout(doc, '重要控制', '涉及健康、过敏、联系方式等敏感信息时，系统只在取得必要授权和明确业务用途的前提下保存；应提供查询、更正和删除入口，并设置保留期限。', color='FDECEC', label_color=RED)

    add_heading(doc, '5. AI 决策、工具调用与人工接管', 1)
    add_text(doc, 'AI 的输出分为“直接回答”“调用工具后回答”“转人工”三种。工具调用必须经过参数校验、权限校验和审计记录，避免由自然语言直接触发不可逆业务操作。')
    add_table(doc, ['场景', 'AI 行为', '控制要求'], [
        ('门店信息与常见项目咨询', '检索知识库后自动回答', '引用门店已发布资料；缺少依据时说明无法确认'),
        ('项目推荐', '追问需求并给出有限推荐', '不做疾病诊断、绝对效果承诺或医疗建议'),
        ('预约意向', '收集项目、时间、联系方式，创建待确认待办', '预约状态默认“待门店确认”，不得伪造已预约成功'),
        ('订单/物流查询', '调用已授权的订单查询工具', '仅查询与已验证客户身份关联的数据'),
        ('投诉、退款、纠纷', '生成摘要并转人工', '不得自行承诺赔付、退款或改价'),
        ('过敏、伤口、孕期、医疗美容', '使用安全提示并转人工或专业人员', '禁止诊断、处方、风险淡化和疗效保证'),
    ], [1730, 3830, 3800])
    add_heading(doc, '5.1 人工接管体验', 2)
    add_text(doc, '转人工时，系统向店员展示对话摘要、客户标签、已提取记忆、客户诉求、推荐回复和转接原因。人工处理后，店员可将有效答案一键保存为门店 FAQ 候选项，形成可审查的知识迭代。')

    add_heading(doc, '6. 数据模型与多租户隔离', 1)
    add_table(doc, ['实体', '核心字段', '用途'], [
        ('Tenant / Store', 'tenant_id、门店资料、行业模板、模型配置、渠道配置', '保证每家门店的数据、知识和规则隔离'),
        ('Customer', 'customer_id、身份标识、授权状态、标签、联系信息', '建立跨会话客户档案'),
        ('Conversation / Message', '会话 ID、渠道、消息内容、状态、处理轨迹', '客服上下文、审计与质检'),
        ('Memory', '主体、类型、内容、置信度、来源、审核状态、过期时间', '可检索、可审核、可删除的长期记忆'),
        ('Knowledge Document', '版本、分段、适用门店、发布日期、来源', '可追溯知识库'),
        ('Task / Appointment Lead', '客户、服务、时间、状态、负责人、跟进记录', '将咨询转为门店可执行待办'),
        ('Audit Log', '工具调用、操作者、参数摘要、结果、时间', '权限追踪、排障和风险审计'),
    ], [1880, 4090, 3390])
    add_text(doc, '数据隔离要求：所有业务表、检索过滤条件和对象存储路径必须携带 tenant_id；后台权限至少区分平台管理员、门店管理员和门店员工。不得在模型检索中跨门店读取数据。', size=10.5, color=MUTED, after=8)

    add_heading(doc, '7. 安全、合规与运行保障', 1)
    add_table(doc, ['风险点', '首期控制措施'], [
        ('不实或越权回答', '知识优先检索、置信度阈值、禁答清单、敏感词与人工接管规则。'),
        ('美容/医美风险', '禁止诊断、疗效保证和处方建议；涉及病症、过敏、孕期、创伤或医疗项目直接升级人工。'),
        ('个人信息', '最小化收集；身份绑定前匿名；明确隐私提示；记忆可查询、更正、删除；访问按角色授权。'),
        ('模型与密钥安全', '服务端保管 API Key；HTTPS；密钥加密；不在客户端、日志和提示词中暴露敏感密钥。'),
        ('服务稳定性', '超时降级为人工留言；任务队列重试；日志、健康检查、每日备份和错误告警。'),
        ('渠道合规', '优先官方 API；不以个人微信或浏览器模拟操作作为生产核心路径；企业微信接入以账户资质和官方可用能力为准。'),
    ], [2250, 7110])

    add_heading(doc, '8. 七日实施计划与验收标准', 1)
    add_table(doc, ['日期', '主要工作', '可验收产出'], [
        ('第 1 天', '项目骨架、账户与门店模型、行业模板、部署环境', '可创建试点门店，完成基础配置与 HTTPS 测试环境'),
        ('第 2 天', '知识库上传、文档解析、检索问答、禁答规则', '能基于门店资料回答 20 个标准问题'),
        ('第 3 天', '客户档案、对话存储、候选记忆提取与后台编辑', '同一手机号再次进入可恢复已确认记忆'),
        ('第 4 天', '网页聊天组件、会话状态、人工接管和店员通知', '客户可对话，店员可查看并接管会话'),
        ('第 5 天', '预约意向流程、待办、纠错反馈和 FAQ 候选', 'AI 能收集预约信息并生成待确认待办'),
        ('第 6 天', '安全控制、日志、监控、备份、企业微信适配预研', '敏感场景按规则转人工；运行日志可追溯'),
        ('第 7 天', '真实资料回归测试、话术调优、部署试运行', '试点门店可通过二维码/链接让真实客户使用'),
    ], [1050, 4870, 3440])
    add_heading(doc, '8.1 MVP 验收指标', 2)
    for item in [
        '知识库标准测试集中，已配置的 20 个高频问题回答正确率不低于 85%；无法确认的信息必须明确说明并转人工。',
        '客户可在 60 秒内完成一次标准咨询；AI 可在一次会话内收集预约所需最小信息。',
        '人工接管时可看到完整会话、客户摘要、转人工原因和 AI 已执行动作。',
        '门店管理员能够编辑资料、发布 FAQ、管理客户记忆，并查看工具调用审计记录。',
        '服务异常时不丢失客户留言，能生成待处理会话或人工跟进任务。',
    ]:
        add_bullet(doc, item)

    add_heading(doc, '9. 后续演进路线', 1)
    add_table(doc, ['阶段', '时间建议', '演进目标'], [
        ('试点优化', '上线后 2 至 4 周', '完善美容行业模板，积累高频问题，调优转人工和预约转化流程。'),
        ('业务闭环', '第 2 阶段', '接入小程序预约、会员、支付、订单、库存和物流查询工具。'),
        ('渠道扩展', '第 3 阶段', '按官方能力接入公众号、企业微信、视频号/微信小店；小红书等渠道单独评估。'),
        ('产品化', '验证 3 至 5 家门店后', '建设行业模板市场、标准化入驻流程、数据看板、计费和多门店运营能力。'),
    ], [1800, 2500, 5060])
    add_callout(doc, '推荐决策', '以“单行业 + 单试点门店 + 网页客服”启动，先验证咨询自动化率、预约转化和人工节省时间。企业微信和交易动作在取得接口能力、账户资质及稳定业务规则后逐步上线。', color=LIGHT_BLUE)

    add_heading(doc, '10. 上线前需由试点门店提供的资料', 1)
    for item in [
        '门店名称、地址、营业时间、联系方式、服务区域与服务人员安排。',
        '服务项目、价格、时长、适用人群、注意事项、禁忌和可预约时段。',
        '历史常见问题、人工话术、投诉与售后处理边界、不可承诺内容。',
        '隐私政策和客户信息保存规则；负责确认预约和处理转人工会话的员工名单。',
        '若接入企业微信或订单系统：对应企业账号、应用管理员、已获批准的官方接口权限。',
    ]:
        add_bullet(doc, item)

    doc.core_properties.title = '个体门店 AI 客服技术方案 v1.0'
    doc.core_properties.subject = '一周试点上线方案'
    doc.core_properties.author = 'AI 客服项目组'
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == '__main__':
    main()
